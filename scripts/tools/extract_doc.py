from docx import Document
import sys

doc = Document(r'C:\Users\Nazifa\supremeai\plans\main plan\SupremeAI_Self_Learning_Documentation.docx')

full_text = []
for para in doc.paragraphs:
    full_text.append(para.text)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            full_text.append(cell.text)

# Write to UTF-8 file to avoid encoding issues
with open('extracted_doc.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(full_text))

print("Document extracted successfully")
